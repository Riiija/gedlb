from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "MANUEL_UTILISATEUR_SOFTAPPLI.md"
OUTPUT = ROOT / "Manuel_utilisateur_SoftAppli.docx"
CAPTURES = ROOT / "captures"
LOGO = ROOT.parent.parent / "public" / "softappli-logo.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "64748B"
TABLE_HEADER = "E8EEF5"
TABLE_BORDER = "B8C7D9"
CALLOUT = "F4F6F9"


ACCENTS = {
    r"\becran\b": "ecran",
}


def fr(text: str) -> str:
    """Restore the most common French accents for a polished Word deliverable."""
    replacements = [
        ("generation", "g\u00e9n\u00e9ration"),
        ("Generale", "G\u00e9n\u00e9rale"),
        ("generale", "g\u00e9n\u00e9rale"),
        ("generales", "g\u00e9n\u00e9rales"),
        ("Presentation", "Pr\u00e9sentation"),
        ("presentation", "pr\u00e9sentation"),
        ("matieres", "mati\u00e8res"),
        ("ecran", "\u00e9cran"),
        ("Ecran", "\u00c9cran"),
        ("ecrans", "\u00e9crans"),
        ("Ecrans", "\u00c9crans"),
        ("deconnexion", "d\u00e9connexion"),
        ("Deconnexion", "D\u00e9connexion"),
        ("acces", "acc\u00e8s"),
        ("Acces", "Acc\u00e8s"),
        ("accede", "acc\u00e8de"),
        ("acceder", "acc\u00e9der"),
        ("apres", "apr\u00e8s"),
        ("debut", "d\u00e9but"),
        ("succes", "succ\u00e8s"),
        ("echec", "\u00e9chec"),
        ("saisir", "saisir"),
        ("saisi", "saisi"),
        ("saisie", "saisie"),
        ("necessaire", "n\u00e9cessaire"),
        ("necessaires", "n\u00e9cessaires"),
        ("identite", "identit\u00e9"),
        ("verifie", "v\u00e9rifie"),
        ("verifier", "v\u00e9rifier"),
        ("masque", "masqu\u00e9"),
        ("masquee", "masqu\u00e9e"),
        ("lisible", "lisible"),
        ("associe", "associ\u00e9"),
        ("associes", "associ\u00e9s"),
        ("connecte", "connect\u00e9"),
        ("connectes", "connect\u00e9s"),
        ("metier", "m\u00e9tier"),
        ("referentiel", "r\u00e9f\u00e9rentiel"),
        ("Reference", "R\u00e9f\u00e9rence"),
        ("reference", "r\u00e9f\u00e9rence"),
        ("references", "r\u00e9f\u00e9rences"),
        ("creation", "cr\u00e9ation"),
        ("Creation", "Cr\u00e9ation"),
        ("creer", "cr\u00e9er"),
        ("Creer", "Cr\u00e9er"),
        ("selection", "s\u00e9lection"),
        ("selectionnee", "s\u00e9lectionn\u00e9e"),
        ("selectionner", "s\u00e9lectionner"),
        ("Selectionner", "S\u00e9lectionner"),
        ("activite", "activit\u00e9"),
        ("activites", "activit\u00e9s"),
        ("desactive", "d\u00e9sactiv\u00e9"),
        ("desactiver", "d\u00e9sactiver"),
        ("utilisee", "utilis\u00e9e"),
        ("utilises", "utilis\u00e9s"),
        ("utilisation", "utilisation"),
        ("disponible", "disponible"),
        ("disponibles", "disponibles"),
        ("modifier", "modifier"),
        ("enregistrer", "enregistrer"),
        ("supprimer", "supprimer"),
        ("confirmation", "confirmation"),
        ("donnees", "donn\u00e9es"),
        ("element", "\u00e9l\u00e9ment"),
        ("elements", "\u00e9l\u00e9ments"),
        ("Element", "\u00c9l\u00e9ment"),
        ("Elements", "\u00c9l\u00e9ments"),
        ("etat", "\u00e9tat"),
        ("Etats", "\u00c9tats"),
        ("resume", "r\u00e9sume"),
        ("Résumé", "R\u00e9sum\u00e9"),
        ("deroulant", "d\u00e9roulant"),
        ("defaut", "d\u00e9faut"),
        ("periode", "p\u00e9riode"),
        ("previsionnelle", "pr\u00e9visionnelle"),
        ("geographique", "g\u00e9ographique"),
        ("coherent", "coh\u00e9rent"),
        ("coherente", "coh\u00e9rente"),
        ("tracabilite", "tra\u00e7abilit\u00e9"),
        ("qualite", "qualit\u00e9"),
        ("telecharger", "t\u00e9l\u00e9charger"),
        ("detail", "d\u00e9tail"),
        ("Details", "D\u00e9tails"),
        ("detaille", "d\u00e9taill\u00e9"),
        ("complete", "compl\u00e8te"),
        ("Complete", "Compl\u00e8te"),
        ("role", "r\u00f4le"),
        ("Role", "R\u00f4le"),
        ("Nouveau depot", "Nouveau d\u00e9p\u00f4t"),
        ("Depot", "D\u00e9p\u00f4t"),
        ("depot", "d\u00e9p\u00f4t"),
    ]
    out = text
    for src, dst in replacements:
        out = out.replace(src, dst)
    return out


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Pt(width / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_run_font(run, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def add_rich_text(paragraph, text: str) -> None:
    text = fr(text)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(INK)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("111827")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 24, INK, 0, 8),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name in {"Title", "Heading 1", "Heading 2", "Heading 3"}
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def cover_page(doc: Document) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.55))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Manuel utilisateur SoftAppli")
    set_run_font(run, size=26, bold=True, color=INK)
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Socle applicatif et shell des modules")
    set_run_font(run, size=13, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(22)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, [2600, 6760])
    entries = [
        ("Destinataires", "Utilisateurs finaux et administrateurs SoftAppli"),
        ("Perimetre", "Connexion, accueil, navigation commune, recherche, notifications, administration utilisateurs, projets/sites et licences"),
        ("Version", "1.0"),
        ("Date", "12 juin 2026"),
    ]
    for i, (label, value) in enumerate(entries):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        c1.text = ""
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(c0, TABLE_HEADER)
        set_cell_margins(c0)
        set_cell_margins(c1)
        r0 = c0.paragraphs[0].add_run(label)
        set_run_font(r0, size=10, bold=True, color=INK)
        r1 = c1.paragraphs[0].add_run(value)
        set_run_font(r1, size=10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document genere a partir de la maquette React SoftAppli")
    set_run_font(r, size=10, italic=True, color=MUTED)
    doc.add_page_break()


def add_table_of_contents(doc: Document, headings: list[tuple[int, str]]) -> None:
    h = doc.add_paragraph("Table des matieres", style="Heading 1")
    keep_with_next(h)
    for level, text in headings:
        if level > 2:
            continue
        p = doc.add_paragraph(style="List Bullet" if level == 2 else "Normal")
        p.paragraph_format.left_indent = Inches(0.25 if level == 2 else 0)
        run = p.add_run(fr(text))
        set_run_font(run, size=10.5 if level == 2 else 11, bold=level == 1, color=INK if level == 1 else "334155")
    doc.add_page_break()


def caption_after_image(doc: Document, alt: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(fr(alt))
    set_run_font(run, size=9, italic=True, color=MUTED)


def add_markdown_image(doc: Document, line: str) -> None:
    match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if not match:
        return
    alt, rel = match.groups()
    path = ROOT / rel
    if not path.exists():
        p = doc.add_paragraph()
        p.add_run(f"[Image manquante: {rel}]").bold = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    caption_after_image(doc, alt)


def table_widths_for(headers: list[str]) -> list[int]:
    n = len(headers)
    if n == 2:
        return [2500, 6860]
    if n == 3:
        return [2200, 3600, 3560]
    if n == 4:
        return [1900, 2600, 2200, 2660]
    if n == 5:
        return [1500, 2200, 1900, 1350, 2410]
    return [int(9360 / n)] * n


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    parsed: list[list[str]] = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) < 2:
        return
    headers = parsed[0]
    body = parsed[2:] if re.match(r"^\s*\|?[-: ]+\|", rows[1]) else parsed[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, table_widths_for(headers))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(fr(header))
        set_run_font(r, size=8.5, bold=True, color=INK)
    for row in body:
        tr = table.add_row()
        for i, value in enumerate(row[: len(headers)]):
            cell = tr.cells[i]
            cell.text = ""
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, value)
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(8.3 if len(headers) >= 5 else 9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_paragraph_from_markdown(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped == "---":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        return
    numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if numbered:
        p = doc.add_paragraph(style="List Number")
        add_rich_text(p, numbered.group(2))
        return
    bullet = re.match(r"^[-*]\s+(.*)$", stripped)
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
        add_rich_text(p, bullet.group(1))
        return
    p = doc.add_paragraph()
    add_rich_text(p, stripped)


def parse_headings(lines: list[str]) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []
    for line in lines:
        m = re.match(r"^(#{1,3})\s+(.+)$", line)
        if not m:
            continue
        text = m.group(2).strip()
        if text.lower().startswith("manuel utilisateur"):
            continue
        headings.append((len(m.group(1)), text))
    return headings


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    cover_page(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    headings = parse_headings(lines)
    add_table_of_contents(doc, headings)

    in_table: list[str] = []
    skip_initial_title = True
    for line in lines:
        if in_table and not line.strip().startswith("|"):
            add_markdown_table(doc, in_table)
            in_table = []
        if line.strip().startswith("|"):
            in_table.append(line)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if skip_initial_title and text.lower().startswith("manuel utilisateur"):
                skip_initial_title = False
                continue
            style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            p = doc.add_paragraph(fr(text), style=style)
            keep_with_next(p)
            continue

        if re.match(r"!\[[^\]]*\]\([^)]+\)", line.strip()):
            add_markdown_image(doc, line)
            continue

        add_paragraph_from_markdown(doc, line)

    if in_table:
        add_markdown_table(doc, in_table)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("SoftAppli - Manuel utilisateur - 12 juin 2026")
    set_run_font(r, size=8.5, color=MUTED)

    doc.core_properties.title = "Manuel utilisateur SoftAppli"
    doc.core_properties.subject = "Socle applicatif et shell des modules"
    doc.core_properties.author = "SoftWell / Documentation utilisateur"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
