from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "MANUEL_UTILISATEUR_SOFTDOCS.md"
OUTPUT = ROOT / "Manuel_utilisateur_SoftDocs.docx"
LOGO = ROOT.parent.parent / "public" / "softdocs-logo.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "64748B"
TABLE_HEADER = "E8EEF5"
TABLE_BORDER = "B8C7D9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_run_font(run, size=None, bold=None, color=None, italic=None, font="Calibri") -> None:
    run.font.name = font
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def add_rich_text(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=9, color=INK, font="Consolas")
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    style_tokens = {
        "Title": (26, INK, 18, 8),
        "Subtitle": (13, MUTED, 0, 22),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in style_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name in {"Title", "Heading 1", "Heading 2", "Heading 3"}
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)


def add_label_value_table(doc: Document, entries: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, [2600, 6760])
    for i, (label, value) in enumerate(entries):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        c1.text = ""
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(c0, TABLE_HEADER)
        set_cell_margins(c0)
        set_cell_margins(c1)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, size=10, bold=True, color=INK)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_rich_text(p1, value)
        for run in p1.runs:
            if run.font.size is None:
                run.font.size = Pt(10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def cover_page(doc: Document) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.7))

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Manuel utilisateur SoftDocs")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Gestion documentaire, validation, rapports et signature SoftSign")

    add_label_value_table(
        doc,
        [
            ("Destinataires", "Utilisateurs finaux, receveurs, validateurs, responsables et administrateurs SoftDocs"),
            ("Périmètre", "Dépôt, réception, validation, refus, redirection, SoftSign, rapports et paramétrage fonctionnel"),
            ("Version", "1.0"),
            ("Date", "12 juin 2026"),
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run("Documentation fonctionnelle rédigée à partir des écrans SoftDocs")
    set_run_font(run, size=10, italic=True, color=MUTED)
    doc.add_page_break()


def add_table_of_contents(doc: Document, headings: list[tuple[int, str]]) -> None:
    h = doc.add_paragraph("Table des matières", style="Heading 1")
    keep_with_next(h)
    for level, text in headings:
        if level > 2:
            continue
        p = doc.add_paragraph(style="List Bullet" if level == 2 else "Normal")
        p.paragraph_format.left_indent = Inches(0.25 if level == 2 else 0)
        run = p.add_run(text)
        set_run_font(run, size=10.5 if level == 2 else 11, bold=level == 1, color=INK if level == 1 else "334155")
    doc.add_page_break()


def caption_after_image(doc: Document, alt: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(alt)
    set_run_font(run, size=9, italic=True, color=MUTED)


def add_markdown_image(doc: Document, line: str) -> None:
    match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if not match:
        return
    alt, rel = match.groups()
    path = ROOT / rel
    if not path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Image manquante : {rel}]")
        set_run_font(run, bold=True, color="9B1C1C")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    caption_after_image(doc, alt)


def table_widths_for(headers: list[str]) -> list[int]:
    n = len(headers)
    if n == 2:
        return [2300, 7060]
    if n == 3:
        return [2200, 3600, 3560]
    if n == 4:
        return [1800, 2700, 2350, 2510]
    if n == 5:
        return [1450, 2200, 1900, 1400, 2410]
    return [int(9360 / n)] * n


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    parsed: list[list[str]] = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) < 2:
        return
    headers = parsed[0]
    body = parsed[2:] if re.match(r"^\s*\|?[-: ]+\|", rows[1]) else parsed[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, table_widths_for(headers))

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=8.7, bold=True, color=INK)

    for row in body:
        tr = table.add_row()
        for i, value in enumerate(row[: len(headers)]):
            cell = tr.cells[i]
            cell.text = ""
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, value)
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(8.7 if len(headers) >= 4 else 9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_paragraph_from_markdown(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped == "---":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        return

    numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if numbered:
        p = doc.add_paragraph(style="List Number")
        add_rich_text(p, numbered.group(2))
        return

    bullet = re.match(r"^[-*]\s+(.*)$", stripped)
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
        add_rich_text(p, bullet.group(1))
        return

    p = doc.add_paragraph()
    add_rich_text(p, stripped)


def parse_headings(lines: list[str]) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []
    for line in lines:
        m = re.match(r"^(#{1,3})\s+(.+)$", line)
        if not m:
            continue
        text = m.group(2).strip()
        if text.lower().startswith("manuel utilisateur"):
            continue
        headings.append((len(m.group(1)), text))
    return headings


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    cover_page(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    add_table_of_contents(doc, parse_headings(lines))

    in_table: list[str] = []
    skip_initial_title = True
    for line in lines:
        if in_table and not line.strip().startswith("|"):
            add_markdown_table(doc, in_table)
            in_table = []
        if line.strip().startswith("|"):
            in_table.append(line)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if skip_initial_title and text.lower().startswith("manuel utilisateur"):
                skip_initial_title = False
                continue
            style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            p = doc.add_paragraph(text, style=style)
            keep_with_next(p)
            continue

        if re.match(r"!\[[^\]]*\]\([^)]+\)", line.strip()):
            add_markdown_image(doc, line)
            continue

        add_paragraph_from_markdown(doc, line)

    if in_table:
        add_markdown_table(doc, in_table)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("SoftDocs - Manuel utilisateur - 12 juin 2026")
    set_run_font(r, size=8.5, color=MUTED)

    doc.core_properties.title = "Manuel utilisateur SoftDocs"
    doc.core_properties.subject = "Gestion documentaire, validation, rapports et signature SoftSign"
    doc.core_properties.author = "SoftWell / Documentation utilisateur"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
