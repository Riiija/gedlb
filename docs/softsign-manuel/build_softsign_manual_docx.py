from __future__ import annotations

import importlib.util
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "MANUEL_UTILISATEUR_SOFTSIGN.md"
OUTPUT = ROOT / "Manuel_utilisateur_SoftSign.docx"
LOGO = ROOT.parent.parent / "public" / "softsign.png"
SOFTDOCS_BUILDER = ROOT.parent / "softdocs-manuel" / "build_softdocs_manual_docx.py"


spec = importlib.util.spec_from_file_location("softdocs_docx_builder", SOFTDOCS_BUILDER)
builder = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(builder)

builder.ROOT = ROOT
builder.SOURCE = SOURCE
builder.OUTPUT = OUTPUT
builder.LOGO = LOGO


def cover_page(doc: Document) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.8))

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Manuel utilisateur SoftSign")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Signature électronique, workflow, OTP, signature externe et suivi des preuves")

    builder.add_label_value_table(
        doc,
        [
            ("Destinataires", "Utilisateurs finaux, déposants, validateurs, signataires, receveurs, fournisseurs et administrateurs SoftSign"),
            ("Périmètre", "Dépôt, workflow, détail document, signature interne, signature externe, OTP, délégation, relances, rapports et paramétrage fonctionnel"),
            ("Version", "1.0"),
            ("Date", "13 juin 2026"),
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run("Documentation fonctionnelle rédigée à partir des écrans SoftSign")
    builder.set_run_font(run, size=10, italic=True, color=builder.MUTED)
    doc.add_page_break()


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    builder.configure_styles(doc)

    cover_page(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    builder.add_table_of_contents(doc, builder.parse_headings(lines))

    in_table: list[str] = []
    skip_initial_title = True
    for line in lines:
        if in_table and not line.strip().startswith("|"):
            builder.add_markdown_table(doc, in_table)
            in_table = []
        if line.strip().startswith("|"):
            in_table.append(line)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if skip_initial_title and text.lower().startswith("manuel utilisateur"):
                skip_initial_title = False
                continue
            style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            p = doc.add_paragraph(text, style=style)
            builder.keep_with_next(p)
            continue

        if re.match(r"!\[[^\]]*\]\([^)]+\)", line.strip()):
            builder.add_markdown_image(doc, line)
            continue

        builder.add_paragraph_from_markdown(doc, line)

    if in_table:
        builder.add_markdown_table(doc, in_table)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SoftSign - Manuel utilisateur - 13 juin 2026")
    builder.set_run_font(run, size=8.5, color=builder.MUTED)

    doc.core_properties.title = "Manuel utilisateur SoftSign"
    doc.core_properties.subject = "Signature électronique, workflow, OTP, signature externe et suivi des preuves"
    doc.core_properties.author = "SoftWell / Documentation utilisateur"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
